"""Markdown 转 Word(docx) 导出工具"""
import re
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_heading_style(paragraph, level: int):
    """设置标题段落样式（颜色 + 字号）"""
    config = {
        1: (18, RGBColor(0x0C, 0x4A, 0x6E)),  # 深蓝
        2: (15, RGBColor(0x0E, 0xA5, 0xE9)),  # 亮蓝
        3: (13, RGBColor(0x02, 0x84, 0xC7)),  # 中蓝
        4: (12, RGBColor(0x07, 0x59, 0x85)),  # 深中蓝
    }
    size, color = config.get(level, (11, RGBColor(0, 0, 0)))
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = True


def _add_horizontal_rule(doc: Document):
    """在文档中插入分割线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_inline(run_text: str, paragraph):
    """解析行内加粗 **text** 并添加到段落"""
    parts = re.split(r"\*\*(.+?)\*\*", run_text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:  # 奇数索引是被 ** 包裹的内容
            run.bold = True


def _parse_table(doc: Document, table_lines: list):
    """解析 Markdown 表格并插入到文档"""
    # 过滤掉分隔行（|---|---|）
    data_rows = [
        line for line in table_lines
        if not re.match(r"^\|[\s\-|:]+\|$", line.strip())
    ]
    if not data_rows:
        return

    rows_data = []
    for line in data_rows:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    col_count = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = re.sub(r"\*\*(.+?)\*\*", r"\1", cell_text)
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # 表格后空一行


def markdown_to_docx(markdown_text: str, output_path: str, title: Optional[str] = None):
    """将 Markdown 文本转换为 Word 文档

    Args:
        markdown_text: Markdown 格式的报告内容
        output_path: 输出 .docx 文件路径
        title: 文档标题（可选，显示在页眉位置）
    """
    doc = Document()

    # 页边距设置（2.5cm 四边）
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    # 默认正文字体
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    lines = markdown_text.splitlines()
    i = 0
    table_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # ---- 跳过 YAML front matter (--- ... ---) ----
        if line.strip() == "---" and i == 0:
            i += 1
            while i < len(lines) and lines[i].strip() != "---":
                i += 1
            i += 1
            continue

        # ---- 表格处理 ----
        if line.strip().startswith("|"):
            table_buffer.append(line)
            in_table = True
            i += 1
            continue
        elif in_table:
            _parse_table(doc, table_buffer)
            table_buffer = []
            in_table = False
            # 不 continue，继续处理当前 line

        # ---- 分割线 ----
        if re.match(r"^-{3,}$", line.strip()):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ---- 标题 ----
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", heading_match.group(2))
            para = doc.add_heading(text, level=level)
            _set_heading_style(para, level)
            i += 1
            continue

        # ---- 无序列表 ----
        bullet_match = re.match(r"^(\s*)-\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            style_name = "List Bullet 2" if indent > 0 else "List Bullet"
            para = doc.add_paragraph(style=style_name)
            _parse_inline(text, para)
            i += 1
            continue

        # ---- 有序列表 ----
        ordered_match = re.match(r"^\d+\.\s+(.*)", line)
        if ordered_match:
            para = doc.add_paragraph(style="List Number")
            _parse_inline(ordered_match.group(1), para)
            i += 1
            continue

        # ---- 空行 ----
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # ---- 普通段落 ----
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        _parse_inline(line.strip(), para)
        i += 1

    # 处理末尾未关闭的表格
    if in_table and table_buffer:
        _parse_table(doc, table_buffer)

    doc.save(output_path)
    return output_path
